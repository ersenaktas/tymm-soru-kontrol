from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as ReportLabImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .review_contract import canonicalize_report_markdown


PACKAGE_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PACKAGE_ROOT / "templates" / "TYMM_Soru_Kontrol_Rapor_Sablonu.docx"
LOGO_PATH = PACKAGE_ROOT / "assets" / "OGM_logo_beyaz_yatay.png"

BLUE = "215E99"
ORANGE = "F1A983"
CYAN = "45B0E1"
GRAY = "808080"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F2F2F2"
PALE_PEACH = "FCE4D6"
PEACH = "F6C5AC"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
DARK_GREEN = "276749"
DARK_TEXT = "17365D"
MUTED_TEXT = "666666"
WHITE = "FFFFFF"
RED = "C00000"
AMBER = "9A6700"

MAIN_HEADINGS = {
    "A — TYMM UYGUNLUĞU",
    "B — BAĞLAM",
    "C — SORU BAZLI DEĞERLENDİRME",
    "D — SET DÜZEYİ DEĞERLENDİRME",
}
METADATA_LABELS = (
    "Ders / Sınıf",
    "Öğrenme Çıktısı",
    "Süreç Bileşeni / Beceri Kodu",
    "Kapsanan Sorular",
    "Genel Sonuç",
)
FIELD_LABELS = (
    "Kapsam",
    "Sonuç",
    "Hata",
    "Hata Açıklaması/Gerekçesi",
    "Kanıt",
    "Düzeltme (Revizyon) Önerisi",
    "ÖNCE",
    "SONRA",
    "Sınırlılık",
    "Gerekli Bilgi",
)
GROUP_LABELS = {
    "soru ve soru cümlesi bulguları": "Soru ve Soru Cümlesi Bulguları",
    "seçenek bulguları": "Seçenek Bulguları",
}


def _normalized(value: str) -> str:
    return " ".join(value.casefold().replace("\ufe0f", "").split())


def _clean_inline(value: str) -> str:
    value = value.replace("\ufe0f", "")
    value = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", value)
    return value.replace("**", "").replace("__", "").replace("`", "").strip()


def _split_label(line: str, labels: tuple[str, ...]) -> tuple[str, str] | None:
    stripped = line.strip()
    for label in sorted(labels, key=len, reverse=True):
        prefix = r"^(?:[-*+]\s+|\d+[.)]\s+)?"
        patterns = (
            prefix + rf"\*\*{re.escape(label)}:\*\*\s*(.*)$",
            prefix + rf"\*\*{re.escape(label)}\*\*\s*:\s*(.*)$",
            prefix + rf"__{re.escape(label)}:__\s*(.*)$",
            prefix + rf"__{re.escape(label)}__\s*:\s*(.*)$",
            prefix + rf"{re.escape(label)}:\s*(.*)$",
        )
        for pattern in patterns:
            match = re.match(pattern, stripped, re.IGNORECASE)
            if match:
                return label, match.group(1).strip()
    return None


def parse_markdown(markdown: str) -> tuple[dict[str, str], list[dict[str, object]]]:
    """Parse the constrained V7 report without changing its decisions."""
    markdown = canonicalize_report_markdown(markdown)
    metadata = {label: "" for label in METADATA_LABELS}
    blocks: list[dict[str, object]] = []
    paragraph: list[str] = []
    pending_empty_field: int | None = None

    def flush() -> None:
        if paragraph:
            blocks.append({"kind": "paragraph", "text": " ".join(paragraph).strip()})
            paragraph.clear()

    for raw in markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            flush()
            pending_empty_field = None
            blocks.append({
                "kind": "heading",
                "level": len(heading.group(1)),
                "text": _clean_inline(heading.group(2)),
            })
            continue
        metadata_line = _split_label(line, METADATA_LABELS)
        if metadata_line:
            flush()
            pending_empty_field = None
            metadata[metadata_line[0]] = _clean_inline(metadata_line[1])
            continue
        field_line = _split_label(line, FIELD_LABELS)
        if field_line:
            flush()
            blocks.append({
                "kind": "field",
                "label": field_line[0],
                "text": _clean_inline(field_line[1]),
            })
            pending_empty_field = len(blocks) - 1 if not field_line[1].strip() else None
            continue
        if pending_empty_field is not None and line.startswith((" ", "\t")):
            continuation = _clean_inline(line.strip())
            if continuation:
                existing = str(blocks[pending_empty_field].get("text", ""))
                blocks[pending_empty_field]["text"] = (existing + " | " + continuation).strip(" |")
                continue
        bold_only = re.match(r"^\*\*(.+?)\*\*\s*$", line.strip())
        if bold_only:
            group_key = _normalized(bold_only.group(1).rstrip(":"))
            if group_key in GROUP_LABELS:
                flush()
                pending_empty_field = None
                blocks.append({"kind": "group", "text": GROUP_LABELS[group_key]})
                continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            flush()
            pending_empty_field = None
            blocks.append({"kind": "bullet", "text": _clean_inline(bullet.group(1))})
            continue
        if line.strip() in {"---", "***"}:
            flush()
            pending_empty_field = None
            continue
        pending_empty_field = None
        paragraph.append(line.strip())
    flush()
    return metadata, blocks


def _status_palette(value: str) -> tuple[str, str]:
    normalized = _normalized(value)
    if "uygun değil" in normalized:
        return PALE_PEACH, RED
    if "düzeltilmeli" in normalized:
        return PALE_YELLOW, AMBER
    if "incelenemedi" in normalized or "sınırlı inceleme" in normalized:
        return PALE_GRAY, MUTED_TEXT
    return PALE_GREEN, DARK_GREEN


def _set_run_font(run, *, size: float = 9.2, color: str = DARK_TEXT, bold: bool = False, italic: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    symbol_font = "Aptos"
    run.font.name = symbol_font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    fonts = run._element.get_or_add_rPr().rFonts
    for script in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(script), symbol_font)


@lru_cache(maxsize=4)
def _status_icon_png(symbol: str) -> bytes:
    """Draw a tiny status badge that Word preserves when exporting to PDF."""
    from PIL import Image, ImageDraw

    scale = 4
    extent = 64
    image = Image.new("RGBA", (extent, extent), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    if symbol == "✅":
        draw.rounded_rectangle((3, 3, 61, 61), radius=13, fill="#2E7D32")
        draw.line((15, 34, 27, 46, 50, 18), fill="white", width=8, joint="curve")
    elif symbol == "❌":
        draw.rounded_rectangle((3, 3, 61, 61), radius=13, fill="#C00000")
        draw.line((18, 18, 46, 46), fill="white", width=8)
        draw.line((46, 18, 18, 46), fill="white", width=8)
    elif symbol == "⚠":
        draw.polygon(((32, 3), (62, 58), (2, 58)), fill="#F2B233")
        draw.line((32, 19, 32, 40), fill="#17365D", width=7)
        draw.ellipse((28, 47, 36, 55), fill="#17365D")
    else:
        draw.ellipse((4, 4, 60, 60), fill="#808080")
        draw.ellipse((17, 17, 47, 47), fill="white")
    destination = BytesIO()
    image.resize((extent // scale, extent // scale), Image.Resampling.LANCZOS).save(destination, format="PNG")
    return destination.getvalue()


def _add_text_with_status_icons(paragraph, value: str, *, size: float, color: str, bold: bool, italic: bool) -> None:
    from docx.shared import Pt

    cursor = 0
    for match in re.finditer(r"✅|❌|⚠(?:\ufe0f)?|⚪", value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor:match.start()].replace("\ufe0f", ""))
            _set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        symbol = match.group(0).replace("\ufe0f", "")
        icon_key = "⚠" if symbol == "⚠" else symbol
        run = paragraph.add_run()
        try:
            run.add_picture(BytesIO(_status_icon_png(icon_key)), height=Pt(size * 0.95))
        except Exception:
            # The pinned teacher package includes Pillow; this text fallback
            # keeps a report readable even in a manually altered environment.
            run.text = {"✅": "[OK]", "❌": "[X]", "⚠": "[!]", "⚪": "[?]"}.get(icon_key, "[?]")
            _set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        cursor = match.end()
    if cursor < len(value):
        run = paragraph.add_run(value[cursor:].replace("\ufe0f", ""))
        _set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def _add_inline_runs(paragraph, text: str, *, size: float = 9.2, color: str = DARK_TEXT, bold: bool = False, italic: bool = False) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|`(.+?)`", text):
        if match.start() > cursor:
            _add_text_with_status_icons(
                paragraph, text[cursor:match.start()], size=size, color=color, bold=bold, italic=italic
            )
        token = (match.group(1) or match.group(2)).replace("\ufe0f", "")
        _add_text_with_status_icons(paragraph, token, size=size, color=color, bold=True, italic=italic)
        cursor = match.end()
    if cursor < len(text):
        _add_text_with_status_icons(paragraph, text[cursor:], size=size, color=color, bold=bold, italic=italic)
    if not text:
        run = paragraph.add_run("")
        _set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def _get_or_add(parent, tag: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    return element


def _shade(container, fill: str) -> None:
    from docx.oxml.ns import qn

    props = container._p.get_or_add_pPr() if hasattr(container, "_p") else container._tc.get_or_add_tcPr()
    shading = _get_or_add(props, "w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _left_border(paragraph, color: str = DARK_GREEN, size: str = "18", space: str = "8") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = _get_or_add(paragraph._p.get_or_add_pPr(), "w:pBdr")
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    for name, value in (("w:val", "single"), ("w:sz", size), ("w:space", space), ("w:color", color)):
        left.set(qn(name), value)


def _cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    margins = _get_or_add(cell._tc.get_or_add_tcPr(), "w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = DARK_TEXT) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    while len(cell.paragraphs) > 1:
        cell._tc.remove(cell.paragraphs[-1]._p)
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    num_pr = paragraph._p.get_or_add_pPr().find(qn("w:numPr"))
    if num_pr is not None:
        paragraph._p.get_or_add_pPr().remove(num_pr)
    _add_inline_runs(paragraph, text, size=9.0, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _add_band(document, text: str, *, style: str, fill: str, color: str, size: float, page_break_before: bool = False):
    from docx.shared import Pt

    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.space_before = Pt(5 if style == "Heading 1" else 3)
    paragraph.paragraph_format.space_after = Pt(5 if style == "Heading 1" else 2)
    paragraph.paragraph_format.left_indent = Pt(4)
    paragraph.paragraph_format.right_indent = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    _shade(paragraph, fill)
    _add_inline_runs(paragraph, text, size=size, color=color, bold=True)
    return paragraph


def _add_field(document, label: str, value: str) -> None:
    from docx.shared import Pt

    fills = {
        "Kapsam": PALE_BLUE,
        "Hata": PEACH,
        "Hata Açıklaması/Gerekçesi": PALE_PEACH,
        "Kanıt": PALE_GRAY,
        "Düzeltme (Revizyon) Önerisi": PALE_GREEN,
        "ÖNCE": PALE_PEACH,
        "SONRA": PALE_GREEN,
        "Sınırlılık": PALE_PEACH,
        "Gerekli Bilgi": PALE_GREEN,
    }
    fill, value_color = fills.get(label, PALE_GRAY), DARK_TEXT
    if label == "Sonuç":
        fill, value_color = _status_palette(value)
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2 if label in {"SONRA", "Gerekli Bilgi", "Düzeltme (Revizyon) Önerisi"} else 0)
    paragraph.paragraph_format.left_indent = Pt(7)
    paragraph.paragraph_format.right_indent = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = label not in {"SONRA", "Gerekli Bilgi", "Düzeltme (Revizyon) Önerisi"}
    _shade(paragraph, fill)
    _left_border(paragraph)
    _add_inline_runs(paragraph, f"{label}: ", size=9.2, bold=True)
    _add_inline_runs(paragraph, value, size=9.2, color=value_color)


def _add_note(document, text: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Pt(7)
    paragraph.paragraph_format.right_indent = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    _shade(paragraph, PALE_GRAY)
    _left_border(paragraph, color="A6A6A6", size="10", space="6")
    _add_inline_runs(paragraph, text, size=8.9, color=MUTED_TEXT, italic=True)


def _restore_template_header_footer(output_path: Path) -> None:
    """Preserve the approved header/footer package parts byte-for-byte."""
    preserved = {
        "word/header1.xml",
        "word/footer1.xml",
        "word/_rels/header1.xml.rels",
        "word/_rels/footer1.xml.rels",
        "word/media/image1.png",
    }
    temporary = output_path.with_name(f"{output_path.stem}.package-tmp.docx")
    with ZipFile(TEMPLATE_PATH, "r") as template_zip, ZipFile(output_path, "r") as output_zip:
        template_names = set(template_zip.namelist())
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as rebuilt_zip:
            for item in output_zip.infolist():
                data = template_zip.read(item.filename) if item.filename in preserved and item.filename in template_names else output_zip.read(item.filename)
                rebuilt_zip.writestr(item, data)
    temporary.replace(output_path)


def create_docx(markdown: str, output_path: Path) -> None:
    """Place a V7 Markdown report into the approved OGM Word template."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f"Kurumsal rapor şablonu bulunamadı: {TEMPLATE_PATH}")
    metadata, blocks = parse_markdown(markdown)
    document = Document(TEMPLATE_PATH)
    if not document.tables or len(document.tables[0].rows) < len(METADATA_LABELS):
        raise ValueError("Kurumsal rapor şablonunun üst bilgi tablosu geçersiz.")

    body = document._element.body
    metadata_table = document.tables[0]
    for child in list(body):
        if child is metadata_table._tbl or child.tag == qn("w:sectPr"):
            continue
        body.remove(child)

    metadata_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata_table.autofit = False
    for index, label in enumerate(METADATA_LABELS):
        row = metadata_table.rows[index]
        row.cells[0].width = Inches(1.75)
        row.cells[1].width = Inches(5.40)
        _set_cell_text(row.cells[0], f"{label}:", bold=True)
        result_color = _status_palette(metadata[label])[1] if label == "Genel Sonuç" else DARK_TEXT
        _set_cell_text(row.cells[1], metadata[label], color=result_color)
        fill = "E7E6E6" if index % 2 == 0 else "F7F7F7"
        _shade(row.cells[0], fill)
        _shade(row.cells[1], fill)

    # A body table as the very first document element can be pulled into the
    # header area by some Word builds during PDF export.  A minimal leading
    # paragraph anchors the first page below the approved template header.
    lead = document.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(0)
    lead.paragraph_format.line_spacing = 0.1
    lead_run = lead.add_run("\u00a0")
    _set_run_font(lead_run, size=1, color=WHITE)
    metadata_table._tbl.addprevious(lead._p)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.3

    current_main = ""
    for block in blocks:
        kind = str(block["kind"])
        text = str(block.get("text", ""))
        if kind == "heading":
            if text == "TYMM SORU KONTROL RAPORU":
                continue
            if text in MAIN_HEADINGS:
                current_main = text[:1]
                # Let Word keep the heading with its following content, but do
                # not force every main section onto a new page. This avoids
                # large blank areas in dense reports.
                _add_band(document, text, style="Heading 1", fill=BLUE, color=WHITE, size=11, page_break_before=False)
                continue
            if re.match(r"^[ABCD]\.(?:\d+(?:\.\d+)?|SET-\d+)\s+[—-]\s+", text, re.IGNORECASE):
                _add_band(document, text, style="Heading 4" if current_main == "C" else "Heading 3", fill=GRAY, color=WHITE, size=9.2)
                continue
            is_question = current_main == "C" and re.match(r"^Soru\s+\d+\s*$", text, re.IGNORECASE) is not None
            _add_band(document, text, style="Heading 2", fill=CYAN if is_question else ORANGE, color=DARK_TEXT, size=10.5)
            continue
        if kind == "group":
            _add_band(document, text, style="Heading 3", fill=ORANGE, color=DARK_TEXT, size=9.7)
        elif kind == "field":
            _add_field(document, str(block.get("label", "")), text)
        elif kind == "bullet":
            # Custom institutional templates may omit Word's optional
            # ``List Bullet`` style.  Keep the report exportable by falling
            # back to the template's guaranteed Normal style and adding the
            # bullet glyph ourselves.
            try:
                paragraph = document.add_paragraph(style="List Bullet")
                bullet_text = text
            except KeyError:
                paragraph = document.add_paragraph(style="Normal")
                paragraph.paragraph_format.left_indent = Pt(12)
                paragraph.paragraph_format.first_line_indent = Pt(-8)
                bullet_text = f"• {text}"
            paragraph.paragraph_format.space_after = Pt(3)
            _add_inline_runs(paragraph, bullet_text, size=9.2)
        elif text in {"Raporlanacak sorun bulunmadı.", "Tek soru bulunduğu için set düzeyi değerlendirme uygulanamaz."}:
            _add_note(document, text)
        elif text:
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_after = Pt(4)
            _add_inline_runs(paragraph, _clean_inline(text), size=9.2)

    document.core_properties.title = "TYMM Soru Kontrol Raporu — V7"
    document.core_properties.subject = "V7 değerlendirmesi; kurumsal Word ve PDF rapor biçimi"
    document.core_properties.keywords = "TYMM, soru kontrol, V7, OGM"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    _restore_template_header_footer(output_path)


def _powershell_literal(value: Path) -> str:
    return "'" + str(value.resolve()).replace("'", "''") + "'"


def _word_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    if os.name != "nt":
        return False
    executable = shutil.which("powershell.exe") or shutil.which("powershell")
    if not executable:
        return False
    source = _powershell_literal(docx_path)
    destination = _powershell_literal(pdf_path)
    script = (
        "$ErrorActionPreference='Stop'; $word=$null; $doc=$null; "
        "try { $word=New-Object -ComObject Word.Application; $word.Visible=$false; "
        "$word.DisplayAlerts=0; $doc=$word.Documents.Open(" + source + ",$false,$true); "
        "$doc.ExportAsFixedFormat(" + destination + ",17); } "
        "finally { if ($doc -ne $null) { $doc.Close(0) }; if ($word -ne $null) { $word.Quit() } }"
    )
    try:
        completed = subprocess.run(
            [executable, "-NoProfile", "-NonInteractive", "-Command", script],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except (OSError, subprocess.TimeoutExpired):
        return False
    return completed.returncode == 0 and pdf_path.is_file() and pdf_path.stat().st_size > 0


def _libreoffice_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        return False
    try:
        completed = subprocess.run(
            [executable, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except (OSError, subprocess.TimeoutExpired):
        return False
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if completed.returncode == 0 and generated.is_file():
        if generated.resolve() != pdf_path.resolve():
            generated.replace(pdf_path)
        return pdf_path.stat().st_size > 0
    return False


def _register_pdf_fonts() -> tuple[str, str]:
    candidates = (
        (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/arialbd.ttf"), "ArialV7"),
        (Path("C:/Windows/Fonts/calibri.ttf"), Path("C:/Windows/Fonts/calibrib.ttf"), "CalibriV7"),
        (Path("C:/Windows/Fonts/segoeui.ttf"), Path("C:/Windows/Fonts/segoeuib.ttf"), "SegoeUIV7"),
        (Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"), Path("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"), "LiberationV7"),
        (Path("/usr/share/fonts/truetype/freefont/FreeSans.ttf"), Path("/usr/share/fonts/truetype/freefont/FreeSansBold.ttf"), "FreeSansV7"),
        (Path("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"), Path("/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf"), "NotoSansV7"),
        (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"), "DejaVuV7"),
    )
    for regular_path, bold_path, family in candidates:
        if regular_path.is_file() and bold_path.is_file():
            try:
                pdfmetrics.registerFont(TTFont(family, str(regular_path)))
                pdfmetrics.registerFont(TTFont(f"{family}-Bold", str(bold_path)))
                return family, f"{family}-Bold"
            except Exception:
                pass

    # Dynamic search in Linux font directories
    font_root = Path("/usr/share/fonts")
    if font_root.is_dir():
        ttfs = list(font_root.rglob("*.ttf"))
        reg = next((f for f in ttfs if "regular" in f.stem.lower() or "sans" in f.stem.lower()), None)
        bold = next((f for f in ttfs if "bold" in f.stem.lower()), None)
        if reg and bold:
            try:
                pdfmetrics.registerFont(TTFont("DynFontV7", str(reg)))
                pdfmetrics.registerFont(TTFont("DynFontV7-Bold", str(bold)))
                return "DynFontV7", "DynFontV7-Bold"
            except Exception:
                pass

    return "Helvetica", "Helvetica-Bold"



def _pdf_text(value: str) -> str:
    replacements = {
        "⚠️ Düzeltilmeli": "[!] DÜZELTİLMELİ",
        "⚠ Düzeltilmeli": "[!] DÜZELTİLMELİ",
        "❌ Uygun Değil": "[X] UYGUN DEĞİL",
        "⚪ İncelenemedi": "[?] İNCELENEMEDİ",
        "⚪ Sınırlı İnceleme": "[?] SINIRLI İNCELEME",
        "✅ Uygun": "[OK] UYGUN",
        "⚠️": "[!]",
        "⚠": "[!]",
        "❌": "[X]",
        "⚪": "[?]",
        "✅": "[OK]",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    return html.escape(_clean_inline(value))


def _fallback_pdf(markdown: str, output_path: Path) -> None:
    metadata, blocks = parse_markdown(markdown)
    regular, bold = _register_pdf_fonts()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.62 * inch,
        rightMargin=0.62 * inch,
        topMargin=1.28 * inch,
        bottomMargin=0.62 * inch,
        title="TYMM Soru Kontrol Raporu — V7",
        author="Ortaöğretim Genel Müdürlüğü",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("V7Body", parent=styles["BodyText"], fontName=regular, fontSize=8.8, leading=11.2, textColor=colors.HexColor(f"#{DARK_TEXT}"), spaceAfter=4)
    note_style = ParagraphStyle("V7Note", parent=body, fontName=regular, fontSize=8.3, leading=10.5, textColor=colors.HexColor(f"#{MUTED_TEXT}"))

    def paragraph(value: str, style=body):
        return Paragraph(_pdf_text(value), style)

    def status_value(value: str, style, width: float):
        match = re.search(r"✅|❌|⚠(?:\ufe0f)?|⚪", value)
        if not match:
            return paragraph(value, style)
        symbol = match.group(0).replace("\ufe0f", "")
        cleaned = (value[:match.start()] + value[match.end():]).strip()
        icon = ReportLabImage(BytesIO(_status_icon_png(symbol)), width=9, height=9)
        text_flowable = Paragraph(html.escape(_clean_inline(cleaned)), style)
        nested = Table([[icon, text_flowable]], colWidths=[13, max(width - 13, 20)])
        nested.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return nested

    def band(value: str, fill: str, foreground: str, size: float):
        style = ParagraphStyle("V7Band", parent=body, fontName=bold, fontSize=size, leading=size + 2, textColor=colors.HexColor(f"#{foreground}"))
        table = Table([[paragraph(value, style)]], colWidths=[document.width], hAlign="LEFT")
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return table

    rows = []
    for label in METADATA_LABELS:
        label_style = ParagraphStyle("V7MetaLabel", parent=body, fontName=bold, fontSize=8.3, leading=10.5)
        value_style = ParagraphStyle("V7MetaValue", parent=body, fontName=regular, fontSize=8.3, leading=10.5)
        meta_value = (
            status_value(metadata[label], value_style, document.width - 1.63 * inch)
            if label == "Genel Sonuç"
            else paragraph(metadata[label], value_style)
        )
        rows.append([paragraph(f"{label}:", label_style), meta_value])
    meta = Table(rows, colWidths=[1.63 * inch, document.width - 1.63 * inch], hAlign="LEFT")
    meta_commands = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D0D0D0")),
    ]
    for index in range(len(rows)):
        meta_commands.append(("BACKGROUND", (0, index), (-1, index), colors.HexColor("#E7E6E6" if index % 2 == 0 else "#F7F7F7")))
    meta.setStyle(TableStyle(meta_commands))
    story = [meta, Spacer(1, 7)]

    current_main = ""
    for block in blocks:
        kind = str(block["kind"])
        text = str(block.get("text", ""))
        if kind == "heading":
            if text == "TYMM SORU KONTROL RAPORU":
                continue
            if text in MAIN_HEADINGS:
                current_main = text[:1]
                # Main sections flow continuously; ReportLab will paginate
                # naturally when the next block no longer fits.
                story.extend([band(text, BLUE, WHITE, 10.5), Spacer(1, 5)])
            elif re.match(r"^[ABCD]\.(?:\d+(?:\.\d+)?|SET-\d+)\s+[—-]\s+", text, re.IGNORECASE):
                story.extend([band(text, GRAY, WHITE, 8.6), Spacer(1, 1.5)])
            else:
                is_question = current_main == "C" and re.match(r"^Soru\s+\d+\s*$", text, re.IGNORECASE)
                story.extend([band(text, CYAN if is_question else ORANGE, DARK_TEXT, 9.5), Spacer(1, 2.5)])
        elif kind == "group":
            story.extend([band(text, ORANGE, DARK_TEXT, 8.9), Spacer(1, 1.5)])
        elif kind == "field":
            label = str(block.get("label", ""))
            fill = {
                "Kapsam": PALE_BLUE,
                "Hata": PEACH,
                "Hata Açıklaması/Gerekçesi": PALE_PEACH,
                "Kanıt": PALE_GRAY,
                "Düzeltme (Revizyon) Önerisi": PALE_GREEN,
                "ÖNCE": PALE_PEACH,
                "SONRA": PALE_GREEN,
                "Sınırlılık": PALE_PEACH,
                "Gerekli Bilgi": PALE_GREEN,
            }.get(label, PALE_GRAY)
            value_color = DARK_TEXT
            if label == "Sonuç":
                fill, value_color = _status_palette(text)
            label_style = ParagraphStyle("V7FieldLabel", parent=body, fontName=bold, fontSize=8.5, leading=10.8)
            value_style = ParagraphStyle("V7FieldValue", parent=body, fontName=regular, fontSize=8.5, leading=10.8, textColor=colors.HexColor(f"#{value_color}"))
            rendered_value = (
                status_value(text, value_style, document.width - 1.63 * inch)
                if label == "Sonuç"
                else paragraph(text, value_style)
            )
            field_table = Table([[paragraph(f"{label}:", label_style), rendered_value]], colWidths=[1.63 * inch, document.width - 1.63 * inch], hAlign="LEFT")
            field_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")),
                ("LINEBEFORE", (0, 0), (0, 0), 2.2, colors.HexColor(f"#{DARK_GREEN}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(field_table)
        elif kind == "bullet":
            story.append(Paragraph("• " + _pdf_text(text), body))
        elif text in {"Raporlanacak sorun bulunmadı.", "Tek soru bulunduğu için set düzeyi değerlendirme uygulanamaz."}:
            note = Table([[paragraph(text, note_style)]], colWidths=[document.width])
            note.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{PALE_GRAY}")),
                ("LINEBEFORE", (0, 0), (0, 0), 1.3, colors.HexColor("#A6A6A6")),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(note)
        elif text:
            story.append(paragraph(text))

    def decorate(canvas, doc):
        width, height = A4
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(f"#{BLUE}"))
        canvas.rect(0, height - 0.98 * inch, width, 0.98 * inch, fill=1, stroke=0)
        if LOGO_PATH.is_file():
            canvas.drawImage(str(LOGO_PATH), 0.48 * inch, height - 0.93 * inch, width=1.85 * inch, height=0.82 * inch, preserveAspectRatio=True, anchor="sw", mask="auto")
        canvas.setFillColor(colors.white)
        canvas.setFont(bold, 11)
        canvas.drawRightString(width - 0.55 * inch, height - 0.47 * inch, "TYMM SORU KONTROL RAPORU")
        canvas.setStrokeColor(colors.HexColor("#D9E2EC"))
        canvas.line(0.62 * inch, 0.44 * inch, width - 0.62 * inch, 0.44 * inch)
        canvas.setFillColor(colors.HexColor(f"#{MUTED_TEXT}"))
        canvas.setFont(regular, 7.5)
        canvas.drawString(0.62 * inch, 0.27 * inch, "Ortaöğretim Genel Müdürlüğü | V7")
        canvas.drawRightString(width - 0.62 * inch, 0.27 * inch, f"Sayfa {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=decorate, onLaterPages=decorate)


def create_pdf(markdown: str, output_path: Path, *, docx_path: Path | None = None) -> None:
    """Create PDF from the branded Word document, with a branded fallback."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    source_docx = docx_path
    temporary_docx: Path | None = None
    if source_docx is None:
        temporary_docx = output_path.with_suffix(".pdf-source.docx")
        create_docx(markdown, temporary_docx)
        source_docx = temporary_docx
    try:
        if _word_to_pdf(source_docx, output_path) or _libreoffice_to_pdf(source_docx, output_path):
            return
        _fallback_pdf(markdown, output_path)
    finally:
        if temporary_docx is not None:
            temporary_docx.unlink(missing_ok=True)


def export_report(markdown: str, stem: str, output_dir: Path) -> tuple[Path, Path]:
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"
    create_docx(markdown, docx_path)
    create_pdf(markdown, pdf_path, docx_path=docx_path)
    return docx_path, pdf_path
