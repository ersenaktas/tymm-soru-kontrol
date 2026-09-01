"""Local subject detection and per-job source resolution."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import unicodedata

from .subject_sources import SubjectSourceError, cleanup_subject_source, stage_remote_subject


SUBJECT_LABELS = {
    "biyoloji": "Biyoloji",
    "fizik": "Fizik",
    "kimya": "Kimya",
    "matematik": "Matematik",
    "cografya": "Coğrafya",
    "tarih": "Tarih",
    "felsefe": "Felsefe",
    "edebiyat": "Türk Dili ve Edebiyatı",
}


@dataclass(frozen=True)
class SubjectResolution:
    key: str
    label: str
    source_path: Path | None
    temporary: bool = False


def normalize(value: str) -> str:
    value = value.casefold().translate(str.maketrans("çğıöşü", "cgiosu"))
    return "".join(char for char in unicodedata.normalize("NFKD", value) if not unicodedata.combining(char))


def extract_document_text(path: Path, pages: int = 8) -> str:
    if path.suffix.lower() in {".md", ".txt"}:
        return path.read_text(encoding="utf-8-sig", errors="replace")[:100_000]
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages[:pages])
    if path.suffix.lower() == ".docx":
        from docx import Document
        document = Document(str(path))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.extend(cell.text for cell in row.cells)
        return "\n".join(chunks)[:100_000]
    return ""


def detect_subject(path: Path) -> str | None:
    filename = normalize(path.stem)
    code_patterns = {
        "biyoloji": r"\bbiy\s*\.\s*\d",
        "fizik": r"\bfiz\s*\.\s*\d",
        "kimya": r"\bkim\s*\.\s*\d",
        "matematik": r"\bmat\s*\.\s*\d",
        "cografya": r"\bcog\s*\.\s*\d",
        "tarih": r"\btar\s*\.\s*\d",
        "felsefe": r"\bfel\s*\.\s*\d",
        "edebiyat": r"\btde\s*\.\s*\d",
    }
    for key, pattern in code_patterns.items():
        if re.search(pattern, filename):
            return key
    content = normalize(extract_document_text(path))
    for key, pattern in code_patterns.items():
        if re.search(pattern, content):
            return key
    phrases = {
        "biyoloji": ("biyoloji dersi", "ders: biyoloji"),
        "fizik": ("fizik dersi", "ders: fizik"),
        "kimya": ("kimya dersi", "ders: kimya"),
        "matematik": ("matematik dersi", "ders: matematik"),
        "cografya": ("cografya dersi", "ders: cografya"),
        "tarih": ("tarih dersi", "ders: tarih"),
        "felsefe": ("felsefe dersi", "ders: felsefe"),
        "edebiyat": ("turk dili ve edebiyati", "edebiyat dersi"),
    }
    matches = [key for key, values in phrases.items() if any(value in content for value in values)]
    return matches[0] if len(matches) == 1 else None


def resolve_subject(
    path: Path,
    requested: str,
    source_path: Path | None,
    configured_sources: dict[str, str],
    root: Path,
    *,
    timeout: int = 20,
) -> SubjectResolution:
    key = requested if requested != "auto" else detect_subject(path)
    if not key:
        raise ValueError(f"Ders güvenle algılanamadı: {path.name}. Arayüzden dersi seçin.")
    if key not in SUBJECT_LABELS:
        raise ValueError(f"Bilinmeyen ders seçimi: {key}")
    candidate = source_path
    temporary = False
    if candidate is None and configured_sources.get(key):
        configured = str(configured_sources[key]).strip()
        if configured.casefold().startswith("https://"):
            try:
                candidate = stage_remote_subject(configured, root, key, timeout=timeout)
                temporary = True
            except SubjectSourceError as exc:
                raise ValueError(f"{SUBJECT_LABELS[key]} ders kaynağı internetten alınamadı: {exc}") from exc
        else:
            candidate = Path(configured)
            if not candidate.is_absolute():
                candidate = root / candidate
    if candidate is not None and not candidate.is_file():
        raise ValueError(f"Ders kaynağı bulunamadı: {candidate}")
    return SubjectResolution(key, SUBJECT_LABELS[key], candidate, temporary)


def discover_question_files(folder: Path) -> list[Path]:
    if not folder.is_dir():
        raise ValueError(f"Klasör bulunamadı: {folder}")
    return sorted(
        [path for path in folder.rglob("*") if path.is_file() and path.suffix.lower() in {".pdf", ".docx"}],
        key=lambda path: str(path).casefold(),
    )
